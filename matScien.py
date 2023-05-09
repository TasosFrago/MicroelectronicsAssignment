from openpyxl import Workbook
from docx import Document
from numpy import log
from math import e
import matplotlib.pyplot as plt
import numpy as np
import itertools

# Temps
temps = [30, 35, 40, 45, 50, 55, 60, 65, 70, 75, 80, 85, 90]
temps = [(i + 273.15) for i in temps]
invTemps = [1/i for i in temps]

# Volts
voltsH = [7.95, 7.65, 6.7, 5.7, 5.0, 4.4, 3.8, 3.1, 2.7, 2.2, 1.8, 1.5, 1.26]

voltsC = [7.6, 7.06, 6.01, 4.98, 4.24, 3.47, 2.95, 2.43, 2.02, 1.72, 1.46, 1.24, 1.08]

s = lambda V : (32 * (10**-3))/ V

# Conductivity
condH = [s(i) for i in voltsH]
condC = [s(i) for i in voltsC]

r = lambda V : V / (16 * (10**-3))
resH = [r(i) for i in voltsH]
resC = [r(i) for i in voltsC]

lcH = [log(i) for i in condH]
lcC = [log(i) for i in condC]

def cal():
    k = 8.625 * (10**-5)

    x = np.array(invTemps)
    y1 = np.array(lcH)
    y2 = np.array(lcC)
    m1, b1 = np.polyfit(x, y1, 1)
    print(m1)
    m2, b2 = np.polyfit(x, y2, 1)
    print(m2)
    mo = (y1[-1]-y1[0])/(x[-1]-x[0])
    print(mo)
    lncon1 = m1*x + b1
    lncon2 = m2*x + b2
    lncon1 = lncon1.tolist()
    lncon2 = lncon2.tolist()
    plt.plot(x, y1, color="orange")
    plt.plot(x, lncon1, color="orange", linestyle="dashed")
    plt.plot(x, y2, color="blue")
    plt.plot(x, lncon2, color="blue", linestyle="dashed")

    plt.xlabel("1/T", fontsize=16)
    plt.ylabel("ln(σ)", fontsize=16)
    plt.legend(['κατα θερμανση', 'προσαρμοσμενη θεωρητικη ευθεια', 'κατα ψυξη', 'προσαρμοσμενη θεωρητικη ευθεια'], loc="lower right")
    plt.show()
    plt.savefig("fig.png")

    for (a, b) in itertools.zip_longest(invTemps, lncon1):
        Eg1 = -2*k*m1
        print(Eg1)
    print("st")
    for (a, b) in itertools.zip_longest(invTemps, lncon2):
        Eg2 = -2*k*m2
        print(Eg2)

def plot() -> None:
    li = [i*-1 for i in lcH]
    x = np.array(invTemps)
    y1 = np.array(li)
    y2 = np.array(lcC)
    m, b = np.polyfit(x, y1, 1)
    liny = m*x +b
    plt.plot(x, y1)
    plt.plot(x, liny)
    plt.show()
    return liny

def word() -> None:
    document = Document()
    table1 = document.add_table(rows=1, cols=5, style='Normal Table')
    row1 = table1.rows[0].cells
    row1[0].text = "Θερμοκρασια δειγματος, Τ, Κ"
    row1[1].text = "1/T, 1/K"
    row1[2].text = "Τάση δειγματος κατα θερμανση, U, Volt"
    row1[3].text = "Αντίσταση κατα θερμανση, R, Ω"
    row1[4].text = "Αγωγιμότητα κατα θερμανση, σ, S/m"

    for (a, b, c, d, e) in itertools.zip_longest(temps, invTemps, voltsH, resH, condH):
        row1 = table1.add_row().cells
        row1[0].text = str(a)
        row1[1].text = str(b)
        row1[2].text = str(c)
        row1[3].text = str(d)
        row1[4].text = str(e)
        print(a, b, c, d, e)

    document.add_paragraph('adding space')

    table2 = document.add_table(rows=1, cols=5, style='Normal Table')
    row2 = table2.rows[0].cells
    row2[0].text = "Θερμοκρασια δειγματος, Τ, Κ"
    row2[1].text = "1/T, 1/K"
    row2[2].text = "Τάση δειγματος κατα ψυξη, U, Volt"
    row2[3].text = "Αντίσταση κατα ψυξη, R, Ω"
    row2[4].text = "Αγωγιμότητα κατα ψυξη, σ, S/m"

    for (a, b, c, d, e) in itertools.zip_longest(temps, invTemps, voltsC, resC, condC):
        row2 = table2.add_row().cells
        row2[0].text = str(a)
        row2[1].text = str(b)
        row2[2].text = str(c)
        row2[3].text = str(d)
        row2[4].text = str(e)
        print(a, b, c, d, e)

    document.save("test.docx")



def sheets(x: list, y1: list, y2: list) -> None:
    excel = "test.xlsx"

    workbook = Workbook()
    sheet = workbook.active

    sheet["A1"] = "1/Θερμοκρασια"
    sheet["B1"] = "Αγωγιμοτητα κατα Θερμανση"
    sheet["C1"] = "Αγωγιμοτητα κατα Ψυξη"

    for i, item in enumerate(x):
        A = "A" + str(i + 2)
        sheet[A] = item

    for i, item in enumerate(y1):
        B = "B" + str(i + 2)
        sheet[B] = item

    for i, item in enumerate(y2):
        C = "C" + str(i + 2)
        sheet[C] = item

    workbook.save(filename=excel)

if __name__=='__main__':
    # word()
    # sheets(invTemps, condH, condC)
    # plot()
    # k = 8.625 * (10**-5)
    # lin = plot()
    # lin = lin.tolist()
    # print(lin)
    # print(type(lin))
    # b = (lin[-1] - lin[0])/(invTemps[-1] - invTemps[0])
    # print(b)
    # for (a, b) in itertools.zip_longest(invTemps, lin):
    #     Eg = 2*k*b
    #     print(Eg)
    # Eg = -2*k*(log(condH[0]) / invTemps[0])
    # print(Eg)
    # Eg = 2*k*b
    # print(Eg)
    # sheets(invTemps, lcH, lcC)
    cal()
